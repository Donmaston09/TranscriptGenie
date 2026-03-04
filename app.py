# transcriptgenie.py
import streamlit as st
from youtube_transcript_api import YouTubeTranscriptApi, TranscriptsDisabled
import whisper
import tempfile
import os
from docx import Document
import yt_dlp

# ----------------------------
# Functions
# ----------------------------

# Extract transcript from YouTube if captions exist
def get_youtube_transcript(video_id):
    try:
        transcript = YouTubeTranscriptApi.get_transcript(video_id)
        lines = []
        for t in transcript:
            start = t['start']
            start_time = f"{int(start//3600):02}:{int((start%3600)//60):02}:{int(start%60):02}"
            lines.append(f"[{start_time}] {t['text']}")
        return "\n".join(lines)
    except TranscriptsDisabled:
        return None
    except Exception as e:
        return f"Error fetching transcript: {e}"

# Download audio from YouTube video (using yt_dlp)
def download_youtube_audio(url):
    try:
        temp_dir = tempfile.mkdtemp()
        output_template = os.path.join(temp_dir, "%(id)s.%(ext)s")

        ydl_opts = {
            "format": "bestaudio/best",
            "outtmpl": output_template,
            "quiet": True,
            "noplaylist": True,
            "postprocessors": [{
                "key": "FFmpegExtractAudio",
                "preferredcodec": "mp3",
                "preferredquality": "128",
            }],
        }

        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(url, download=True)
            audio_path = os.path.join(temp_dir, f"{info['id']}.mp3")

        if not os.path.exists(audio_path):
            raise FileNotFoundError("Audio file was not properly created.")
        return audio_path

    except Exception as e:
        return f"Error downloading audio: {e}"

# Transcribe local file using Whisper
def get_local_transcript(file_path):
    model = whisper.load_model("base")
    result = model.transcribe(file_path)
    transcript = []
    for seg in result["segments"]:
        start, end, text = seg["start"], seg["end"], seg["text"].strip()
        start_time = f"{int(start//3600):02}:{int((start%3600)//60):02}:{int(start%60):02}"
        end_time = f"{int(end//3600):02}:{int((end%3600)//60):02}:{int(end%60):02}"
        transcript.append(f"[{start_time} - {end_time}] {text}")
    return "\n".join(transcript)

# Save transcript to Word file
def save_to_word(transcript_text, filename="transcript.docx"):
    doc = Document()
    doc.add_heading("Video Transcript with Timestamps", 0)
    for line in transcript_text.split("\n"):
        doc.add_paragraph(line)
    doc.save(filename)
    return filename

# ----------------------------
# Streamlit App
# ----------------------------

st.set_page_config(page_title="TranscriptGenie", page_icon="🎙️", layout="centered")
st.title("🎙️ TranscriptGenie")
st.subheader("Generate & download transcripts from YouTube or local videos with timestamps")

st.write("Upload a video file or paste a YouTube link to generate a transcript (with timestamps), then export it as a Word file.")

option = st.radio("Choose input type:", ["YouTube Link", "Upload Local Video"])
transcript_text = ""

if option == "YouTube Link":
    youtube_url = st.text_input("Enter YouTube Video URL:")
    if youtube_url and st.button("Generate Transcript"):
        with st.spinner("Fetching YouTube transcript..."):
            video_id = youtube_url.split("v=")[-1].split("&")[0]
            transcript_text = get_youtube_transcript(video_id)

            if transcript_text is None or transcript_text.startswith("Error"):
                st.warning("⚠️ No captions found or transcript unavailable. Attempting audio transcription...")
                with st.spinner("Downloading and transcribing audio... (this may take a few minutes)"):
                    audio_path = download_youtube_audio(youtube_url)
                    if os.path.exists(audio_path):
                        transcript_text = get_local_transcript(audio_path)
                        os.remove(audio_path)
                    else:
                        transcript_text = f"Error: Could not process video.\n{audio_path}"
        st.text_area("Transcript", transcript_text, height=300)

elif option == "Upload Local Video":
    uploaded_file = st.file_uploader("Upload a video file", type=["mp3", "wav", "m4a", "mp4", "mov", "avi", "mkv"])
    if uploaded_file is not None:
        if st.button("Generate Transcript"):
            with tempfile.NamedTemporaryFile(delete=False, suffix=".mp4") as tmp:
                tmp.write(uploaded_file.read())
                tmp_path = tmp.name
            with st.spinner("Transcribing video with Whisper..."):
                transcript_text = get_local_transcript(tmp_path)
            st.text_area("Transcript", transcript_text, height=300)
            os.remove(tmp_path)

# Download Transcript
if transcript_text:
    if st.button("Download Transcript as Word"):
        file_path = save_to_word(transcript_text)
        with open(file_path, "rb") as f:
            st.download_button("📥 Download Word File", f, file_name="transcript.docx")

# ----------------------------
# Sidebar / Footer
# ----------------------------
st.sidebar.markdown("### 📌 About the Author")
st.sidebar.write("**Anthony Onoja**")
st.sidebar.write("University of Surrey, UK")
st.sidebar.write("📧 donmaston09@gmail.com")
st.sidebar.info("Contact me for AI, Data Science & Educational Tech collaborations.")
